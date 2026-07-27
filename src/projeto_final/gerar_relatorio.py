"""Preenche o modelo oficial e exporta o relatório final em DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


PROJECT_DIR = Path("/home/mauro/Projetos/fastcamp-dados-sinteticos")
TEMPLATE = Path("/home/mauro/Downloads/Modelo_Relatório.docx")
OUTPUT_DIR = PROJECT_DIR / "entregas"
OUTPUT_DOCX = OUTPUT_DIR / "Relatório 12 - Mauro Andrade.docx"


def limpar_corpo(documento):
    corpo = documento._element.body
    for elemento in list(corpo):
        if not elemento.tag.endswith("sectPr"):
            corpo.remove(elemento)


def adicionar_titulo(documento):
    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Relatório 12 - Prática: Projeto Final (V)")
    run.bold = True
    run.font.size = Pt(14)

    aluno = documento.add_paragraph()
    aluno.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = aluno.add_run("Mauro Andrade")
    run.font.size = Pt(12)


def adicionar_secao(documento, titulo):
    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.space_before = Pt(12)
    paragrafo.paragraph_format.space_after = Pt(5)
    run = paragrafo.add_run(titulo)
    run.bold = True
    run.font.size = Pt(12)


def adicionar_texto(documento, texto):
    paragrafo = documento.add_paragraph(texto)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo.paragraph_format.first_line_indent = Cm(1.25)
    paragrafo.paragraph_format.space_after = Pt(6)


def adicionar_imagem(documento, caminho, legenda, largura=Cm(12.8)):
    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.add_run().add_picture(str(caminho), width=largura)
    legenda_p = documento.add_paragraph(legenda)
    legenda_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legenda_p.runs[0].italic = True
    legenda_p.runs[0].font.size = Pt(9)


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Modelo oficial não encontrado: {TEMPLATE}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    documento = Document(TEMPLATE)
    limpar_corpo(documento)

    estilos = documento.styles
    estilos["Normal"].font.name = "Arial"
    estilos["Normal"].font.size = Pt(11)
    secao = documento.sections[0]
    secao.top_margin = Cm(2.5)
    secao.bottom_margin = Cm(2.5)
    secao.left_margin = Cm(3)
    secao.right_margin = Cm(2)

    adicionar_titulo(documento)
    adicionar_secao(documento, "Descrição da atividade")

    adicionar_texto(
        documento,
        "O projeto LitterWatch foi desenvolvido como prova de conceito para "
        "detectar e contar visitas de um gato a uma caixa de areia. A proposta "
        "surgiu da dificuldade de acompanhar esse comportamento manualmente, "
        "principalmente durante a noite ou quando o responsável está ausente. "
        "O sistema não diagnostica doenças e não confirma se houve micção ou "
        "defecação; ele apenas demonstra como dados sintéticos e visão "
        "computacional podem apoiar o registro de presença, início, fim e "
        "duração de visitas.",
    )
    adicionar_texto(
        documento,
        "A cena foi criada programaticamente no Blender 3.6.23. Para manter o "
        "projeto compatível com o computador disponível, foi utilizado um gato "
        "3D simplificado formado por treze componentes geométricos. Também foram "
        "modelados a base, as paredes e a areia da caixa. Um script em Python "
        "com a API bpy configurou câmera, piso, materiais, iluminação e "
        "renderização.",
    )
    adicionar_imagem(
        documento,
        PROJECT_DIR / "outputs" / "projeto_final" / "cena_base.png",
        "Figura 1 — Cena-base criada no Blender com gato e caixa de areia.",
    )

    adicionar_texto(
        documento,
        "O gerador automatizou sessenta imagens sintéticas em resolução de "
        "320 por 320 pixels. Foram criadas vinte cenas com o gato dentro da "
        "caixa, vinte próximo e vinte fora. Posição, rotação, cor do gato, "
        "intensidade e posição da luz e altura da câmera foram variadas. O "
        "dataset foi dividido em 48 imagens de treino, 6 de validação e 6 de "
        "teste, com semente fixa para tornar o experimento reproduzível.",
    )
    adicionar_texto(
        documento,
        "As anotações foram geradas sem marcação manual. Os vértices das caixas "
        "tridimensionais de todas as partes do gato foram projetados na câmera "
        "e reunidos em uma única bounding box no formato YOLO. Um segundo "
        "programa verificou a presença de sessenta pares de imagem e rótulo, "
        "valores normalizados entre zero e um e caixas com área positiva.",
    )
    adicionar_imagem(
        documento,
        PROJECT_DIR / "outputs" / "projeto_final" / "anotacoes" / "exemplo_dentro.png",
        "Figura 2 — Exemplo de anotação automática da classe gato.",
        Cm(10.5),
    )

    adicionar_texto(
        documento,
        "Foi utilizado o modelo YOLO26n por transferência de aprendizado. A "
        "arquitetura nano foi escolhida para reduzir o custo computacional. O "
        "treinamento foi realizado em CPU por oito épocas, com imagens de 320 "
        "pixels, batch 2 e uma única classe, denominada gato. O objetivo foi "
        "validar o pipeline completo e não produzir um sistema clínico ou "
        "comercial.",
    )
    adicionar_texto(
        documento,
        "Na validação sintética, o modelo alcançou mAP50 de 0,942. No conjunto "
        "de teste separado, o mAP50 foi 0,636. A diferença entre os valores "
        "indica que o conjunto pequeno não permite afirmar forte capacidade de "
        "generalização. Ainda assim, as seis imagens da demonstração receberam "
        "detecções, permitindo comprovar o funcionamento técnico do fluxo.",
    )
    adicionar_imagem(
        documento,
        PROJECT_DIR
        / "outputs"
        / "projeto_final"
        / "yolo_runs"
        / "litterwatch_yolo26n"
        / "results.png",
        "Figura 3 — Gráficos produzidos durante as oito épocas de treinamento.",
        Cm(16),
    )

    adicionar_texto(
        documento,
        "Por fim, foi implementada uma máquina de estados para transformar "
        "quadros consecutivos em visitas. O contador exige duas confirmações "
        "para entrada e duas para saída, evitando que uma detecção isolada seja "
        "interpretada como nova visita. Em uma sequência simulada, foram "
        "contabilizadas duas visitas, com durações de cinco e quatro segundos, "
        "e um ruído isolado foi ignorado.",
    )

    adicionar_secao(documento, "Dificuldades")
    adicionar_texto(
        documento,
        "A primeira dificuldade foi a limitação de hardware. A versão mais "
        "recente do Blender apresentou incompatibilidade gráfica, por isso foi "
        "adotado o Blender 3.6.23 e o Eevee com renderização em CPU. O "
        "treinamento também foi configurado com modelo nano, resolução pequena "
        "e batch reduzido.",
    )
    adicionar_texto(
        documento,
        "Outra dificuldade foi representar um gato sem utilizar um ativo 3D "
        "pesado. A solução foi construir um modelo simplificado com formas "
        "geométricas, suficiente para demonstrar scripting, randomização e "
        "anotação. Durante os testes, a contagem esperada das partes estava "
        "incorreta; o teste automático revelou que o modelo possuía treze "
        "componentes e a validação foi corrigida antes da geração do dataset.",
    )
    adicionar_texto(
        documento,
        "Também foi necessário calcular uma única caixa para um objeto composto. "
        "A solução consistiu em projetar todos os vértices das treze partes na "
        "imagem e combinar seus limites. Depois, previews com retângulos "
        "desenhados foram inspecionados visualmente para confirmar o alinhamento.",
    )
    adicionar_texto(
        documento,
        "A principal dificuldade conceitual foi manter o escopo responsável. "
        "Contar visitas pode auxiliar a observação, mas não identifica micção, "
        "volume urinário ou doença. Por isso, diagnóstico, sensores, câmera ao "
        "vivo e múltiplos gatos foram explicitamente retirados do MVP.",
    )

    adicionar_secao(documento, "Conclusões")
    adicionar_texto(
        documento,
        "O trabalho integrou as etapas exigidas no desafio: modelagem e "
        "scripting no Blender, randomização, geração de imagens, anotação "
        "automática, organização do dataset, treinamento, avaliação e "
        "demonstração de uma regra de negócio. O uso de scripts tornou o "
        "processo reproduzível e eliminou a necessidade de anotar manualmente "
        "cada imagem.",
    )
    adicionar_texto(
        documento,
        "Os resultados sintéticos confirmam a viabilidade acadêmica, mas não "
        "comprovam funcionamento no mundo real. O gato simplificado, o conjunto "
        "pequeno e a ausência de imagens reais criam um domain gap significativo. "
        "Uma continuação exigiria dados reais autorizados, avaliação em vídeo, "
        "sensores complementares e participação de profissionais veterinários.",
    )
    adicionar_texto(
        documento,
        "Concluo que dados sintéticos são úteis para construir e testar um "
        "pipeline quando dados reais são escassos, mas suas métricas devem ser "
        "interpretadas com cuidado. O projeto cumpre seu objetivo como prova de "
        "conceito educacional e mantém transparentes suas limitações.",
    )

    adicionar_secao(documento, "Referências")
    referencias = [
        "BLENDER FOUNDATION. Blender Python API Documentation. Disponível em: https://docs.blender.org/api/3.6/.",
        "ULTRALYTICS. Ultralytics YOLO Documentation. Disponível em: https://docs.ultralytics.com/.",
        "DENG, P. et al. Quantification of Urine Elimination Behaviors in Cats with a Video Recording System. Journal of Veterinary Internal Medicine, 2017.",
        "SUNG, Y. et al. Multi-Cat Monitoring System Based on Concept Drift Adaptive Machine Learning Architecture. Sensors, 2023.",
        "EMC/UFG. FastCamp de Dados Sintéticos para Inteligência Artificial e Visão Computacional, 2026.",
        "Código e artefatos do projeto: https://github.com/mhaurinho/litterwatch-dados-sinteticos.",
    ]
    for referencia in referencias:
        paragrafo = documento.add_paragraph(style=None)
        paragrafo.paragraph_format.left_indent = Cm(0.6)
        paragrafo.paragraph_format.first_line_indent = Cm(-0.6)
        paragrafo.add_run(referencia)

    documento.save(OUTPUT_DOCX)
    print(f"Relatório DOCX salvo em: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
